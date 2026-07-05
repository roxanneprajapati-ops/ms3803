import os
import re
import pdfplumber
import docx
from docx import Document
from dotenv import load_dotenv
from google import genai


# ============================================================
# STEP 1: LOAD API KEY
# ============================================================

load_dotenv()

api_key = os.getenv("GEMINI_API_KEY")

if not api_key:
    raise ValueError("GEMINI_API_KEY not found. Please add it to your .env file.")

client = genai.Client(api_key=api_key)


# ============================================================
# STEP 2: PROMPTS
# ============================================================

FEEDBACK_PROMPT = """
You are a professional AI recruiter and career advisor with expertise in IT, Software Engineering,
Data Analytics, and Computer Science.

Analyze the candidate's CV provided below.

Provide:
1. Main professional field
2. Two strongest areas of expertise
3. 2-3 suitable job roles
4. Strengths of the CV
5. Weaknesses of the CV
6. 3 actionable recommendations to improve the CV:
   - content
   - structure
   - presentation
   - ATS keyword optimization

Keep the response concise, professional, and easy to understand.

CV Text:
{cv_text}
"""


OPTIMIZED_CV_PROMPT = """
You are an expert CV writer.

Using the original CV and the feedback below, rewrite the CV into an improved ATS-friendly version.

Rules:
- Do not invent fake jobs, education, certifications, or skills.
- Keep the candidate's real experience.
- Improve clarity, structure, grammar, and professional wording.
- Use strong action verbs.
- Make the CV easy to read and suitable for job applications.
- Use clear sections:
  Profile Summary
  Key Skills
  Professional Experience
  Education
  Projects
  Certifications

Original CV:
{cv_text}

Feedback:
{feedback}
"""


# ============================================================
# STEP 3: EXTRACT TEXT FROM CV
# ============================================================

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    text = ""

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()

            if page_text:
                text += page_text + "\n"

    return text.strip()


def extract_text_from_docx(file_path):
    """Extract text from a DOCX file."""
    document = docx.Document(file_path)
    text = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text.strip())

    return "\n".join(text)


def extract_cv_text(file_path):
    """Check file type and extract CV text."""
    if file_path.lower().endswith(".pdf"):
        return extract_text_from_pdf(file_path)

    elif file_path.lower().endswith(".docx"):
        return extract_text_from_docx(file_path)

    else:
        raise ValueError("Unsupported file format. Please use PDF or DOCX.")


# ============================================================
# STEP 4: AI CV ANALYSIS
# ============================================================

def analyze_cv(cv_text):
    """Generate feedback for the CV using Gemini."""
    prompt = FEEDBACK_PROMPT.format(cv_text=cv_text)

    response = client.models.generate_content(
        model="gemini-3-flash-preview",
        contents=prompt,
        config={
                "temperature": 0.7,
                "top_p": 0.9,
                "top_k": 40,
                "max_output_tokens": 2000,
        }
    )

    return response.text


def generate_optimized_cv(cv_text, feedback):
    """Generate improved CV text using Gemini."""
    prompt = OPTIMIZED_CV_PROMPT.format(
        cv_text=cv_text,
        feedback=feedback
    )

    response = client.models.generate_content(
        model="gemini-3-flash-preview",
        contents=prompt
    )

    return response.text


# ============================================================
# STEP 5: CLEAN AI OUTPUT
# ============================================================

def clean_markdown(text):
    """Clean simple markdown formatting from AI response."""
    text = re.sub(r'^\s*\*\s+', '- ', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'#{1,6}\s*', '', text)
    text = re.sub(r'-\s+', '- ', text)

    return text.strip()


# ============================================================
# STEP 6: SAVE OPTIMIZED CV AS DOCX
# ============================================================

def save_text_file(text, output_file):
    """Save feedback or optimized CV as a text file."""
    with open(output_file, "w", encoding="utf-8") as file:
        file.write(text)


def generate_cv_docx(optimized_text, output_file="optimized_cv.docx"):
    """Save the optimized CV as a DOCX file."""
    document = Document()

    document.add_heading("Optimized CV", level=1)

    for line in optimized_text.split("\n"):
        line = line.strip()

        if not line:
            continue

        # Make common CV section titles headings
        clean_line = line.replace(":", "").strip()

        section_titles = [
            "Profile Summary",
            "Key Skills",
            "Professional Experience",
            "Education",
            "Projects",
            "Certifications",
            "Technical Skills",
            "Work Experience"
        ]

        if clean_line in section_titles:
            document.add_heading(clean_line, level=2)
        elif line.startswith("-"):
            document.add_paragraph(line, style="List Bullet")
        else:
            document.add_paragraph(line)

    document.save(output_file)


# ============================================================
# STEP 7: MAIN PROGRAM
# ============================================================

def main():
    print("=" * 60)
    print("AI CV FEEDBACK AND OPTIMIZATION TOOL")
    print("=" * 60)

    if os.path.exists("resume-sample.pdf"):
        file_path = "resume-sample.pdf"
    elif os.path.exists("resume-sample.docx"):
        file_path = "resume-sample.docx"
    else:
        print("No resume found.")
        print("Please place resume.pdf or resume.docx in the same folder as main.py.")
        return



    try:
        print("\nExtracting CV text...")
        cv_text = extract_cv_text(file_path)

        if not cv_text:
            print("No text found in the CV file.")
            return

        print("\nAnalyzing CV with Google Gemini...")
        feedback = analyze_cv(cv_text)
        cleaned_feedback = clean_markdown(feedback)

        print("\n" + "=" * 60)
        print("CV ANALYSIS RESULTS")
        print("=" * 60)
        print(cleaned_feedback)

        save_text_file(cleaned_feedback, "cv_feedback.txt")
        print("\nFeedback saved as: cv_feedback.txt")

        print("\nGenerating optimized CV...")
        optimized_cv = generate_optimized_cv(cv_text, cleaned_feedback)
        cleaned_optimized_cv = clean_markdown(optimized_cv)

        save_text_file(cleaned_optimized_cv, "optimized_cv.txt")
        generate_cv_docx(cleaned_optimized_cv, "optimized_cv.docx")

        print("\n" + "=" * 60)
        print("OPTIMIZED CV GENERATED SUCCESSFULLY")
        print("=" * 60)
        print("Saved files:")
        print("- cv_feedback.txt")
        print("- optimized_cv.txt")
        print("- optimized_cv.docx")

    except Exception as error:
        print("\nAn error occurred:")
        print(error)


if __name__ == "__main__":
    main()